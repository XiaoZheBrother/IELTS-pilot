"""Build the v0.9 product manual as a styled DOCX from the tracked Markdown source."""

from __future__ import annotations

import argparse
import re
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate, Frame, Image as PdfImage, KeepTogether, ListFlowable, ListItem,
    PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle, XPreformatted,
)


INK = "17201B"
MUTED = "59625D"
ACCENT = "245BFF"
HEADING = "2E74B5"
HEADING_DARK = "1F4D78"
LINE = "C8CCC8"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CODE_FILL = "EEF0ED"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), LINE)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def add_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("IELTS Pilot v0.9  |  ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, HEADING, 18, 10),
        "Heading 2": (13, HEADING, 14, 7),
        "Heading 3": (12, HEADING_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("IELTS PILOT  /  PRODUCT GUIDE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])

    doc.core_properties.title = "IELTS Pilot v0.9 产品功能说明书"
    doc.core_properties.subject = "IELTS Pilot v0.9 产品、部署、评分与安全说明"
    doc.core_properties.author = "XiaoZheBrother"
    doc.core_properties.keywords = "IELTS Pilot, reading, writing, AI, product guide"


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    level.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline(paragraph, text, *, size=None, color=INK):
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=size or 9.5, color=HEADING_DARK)
        else:
            label, url = re.match(r"\[(.+?)\]\((.+?)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_run_font(run, size=size, color=HEADING, bold=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def set_paragraph_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(86)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("PRODUCT GUIDE  /  RELEASE 0.9")
    set_run_font(run, size=10, color=ACCENT, bold=True)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IELTS Pilot v0.9")
    set_run_font(run, size=30, color=INK, bold=True)
    title.paragraph_format.space_after = Pt(9)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("产品功能说明书")
    set_run_font(run, size=20, color=HEADING_DARK, bold=True)
    subtitle.paragraph_format.space_after = Pt(12)
    deck = doc.add_paragraph()
    deck.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = deck.add_run("本地优先的 IELTS 阅读与 AI 写作练习工作台")
    set_run_font(run, size=12.5, color=MUTED)
    deck.paragraph_format.space_after = Pt(58)
    for line in ("版本 0.9.0", "2026-08-12", "浏览器版 · Windows 桌面版 · NSIS 安装包"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, line, size=10, color=MUTED)
        p.paragraph_format.space_after = Pt(4)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(48)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.left_indent = Inches(0.55)
    note.paragraph_format.right_indent = Inches(0.55)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(note, "非官方 IELTS 产品 · 内置材料为项目原创 · 所有 Band 仅用于学习反馈", size=9, color=MUTED)


def column_widths(rows):
    columns = len(rows[0])
    maxima = []
    for index in range(columns):
        maxima.append(max(4, min(32, max(len(re.sub(r"[`*]", "", row[index])) for row in rows))))
    if columns and all(row[0].strip().isdigit() or row[0].strip() in {"领域", "入口", "任务", "路由", "场景", "编号"} for row in rows):
        maxima[0] = min(maxima[0], 7)
    weights = [max(0.75, value) for value in maxima]
    total = sum(weights)
    widths = [round(CONTENT_DXA * value / total) for value in weights]
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, column_widths(rows))
    add_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if column_index == 0 and len(value) < 16:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, value, size=9.2, color=INK)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def add_image(doc, source_dir, alt, target):
    path = (source_dir / target).resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.3))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    paragraph.paragraph_format.space_after = Pt(2)
    caption = doc.add_paragraph(style="Caption")
    caption.add_run(f"图：{alt}")


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(paragraph, text, size=10, color=HEADING_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    paragraph_properties.append(shading)
    paragraph.paragraph_format.left_indent = Inches(0.1)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.7, color=INK)


def parse_manual(doc, markdown_path):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    first_break = lines.index("<!-- pagebreak -->")
    lines = lines[first_break + 1:]
    index = 0
    list_kind = None
    list_num_id = None
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            list_kind = None
            list_num_id = None
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            flush_paragraph()
            doc.add_page_break()
            list_kind = None
            list_num_id = None
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            add_code(doc, code)
            index += 1
            continue
        image = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image:
            flush_paragraph()
            add_image(doc, markdown_path.parent, image.group(1), image.group(2))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.fullmatch(r"\|?[\s:|-]+\|?", lines[index + 1].strip()):
            flush_paragraph()
            rows = []
            rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = min(3, len(heading.group(1)))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2))
            index += 1
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            add_callout(doc, stripped[2:])
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or ordered:
            flush_paragraph()
            kind = "bullet" if bullet else "decimal"
            if list_kind != kind or list_num_id is None:
                list_kind = kind
                list_num_id = add_numbering(doc, kind)
            paragraph = doc.add_paragraph()
            set_paragraph_numbering(paragraph, list_num_id)
            add_inline(paragraph, (bullet or ordered).group(1))
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


def register_pdf_fonts():
    pdfmetrics.registerFont(TTFont("MicrosoftYaHei", r"C:\Windows\Fonts\msyh.ttc", subfontIndex=0))
    pdfmetrics.registerFont(TTFont("MicrosoftYaHei-Bold", r"C:\Windows\Fonts\msyhbd.ttc", subfontIndex=0))
    pdfmetrics.registerFont(TTFont("Consolas", r"C:\Windows\Fonts\consola.ttf"))
    pdfmetrics.registerFont(TTFont("Consolas-Bold", r"C:\Windows\Fonts\consolab.ttf"))
    pdfmetrics.registerFontFamily("MicrosoftYaHei", normal="MicrosoftYaHei", bold="MicrosoftYaHei-Bold", italic="MicrosoftYaHei", boldItalic="MicrosoftYaHei-Bold")
    pdfmetrics.registerFontFamily("Consolas", normal="Consolas", bold="Consolas-Bold", italic="Consolas", boldItalic="Consolas-Bold")


def pdf_styles():
    return {
        "body": ParagraphStyle("Body", fontName="MicrosoftYaHei", fontSize=9.6, leading=14.5, textColor=colors.HexColor(f"#{INK}"), spaceAfter=6),
        "h1": ParagraphStyle("H1", fontName="MicrosoftYaHei-Bold", fontSize=16, leading=21, textColor=colors.HexColor(f"#{HEADING}"), spaceBefore=18, spaceAfter=10, keepWithNext=1),
        "h2": ParagraphStyle("H2", fontName="MicrosoftYaHei-Bold", fontSize=13, leading=18, textColor=colors.HexColor(f"#{HEADING}"), spaceBefore=14, spaceAfter=7, keepWithNext=1),
        "h3": ParagraphStyle("H3", fontName="MicrosoftYaHei-Bold", fontSize=11.5, leading=16, textColor=colors.HexColor(f"#{HEADING_DARK}"), spaceBefore=10, spaceAfter=5, keepWithNext=1),
        "caption": ParagraphStyle("Caption", fontName="MicrosoftYaHei", fontSize=8, leading=11, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceBefore=3, spaceAfter=8),
        "cover_kicker": ParagraphStyle("CoverKicker", fontName="MicrosoftYaHei-Bold", fontSize=10, leading=14, textColor=colors.HexColor(f"#{ACCENT}"), alignment=TA_CENTER, spaceAfter=18),
        "cover_title": ParagraphStyle("CoverTitle", fontName="MicrosoftYaHei-Bold", fontSize=29, leading=36, textColor=colors.HexColor(f"#{INK}"), alignment=TA_CENTER, spaceAfter=8),
        "cover_subtitle": ParagraphStyle("CoverSubtitle", fontName="MicrosoftYaHei-Bold", fontSize=19, leading=25, textColor=colors.HexColor(f"#{HEADING_DARK}"), alignment=TA_CENTER, spaceAfter=10),
        "cover_deck": ParagraphStyle("CoverDeck", fontName="MicrosoftYaHei", fontSize=12, leading=18, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceAfter=46),
        "cover_meta": ParagraphStyle("CoverMeta", fontName="MicrosoftYaHei", fontSize=9, leading=14, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceAfter=3),
        "code": ParagraphStyle("Code", fontName="MicrosoftYaHei", fontSize=7.8, leading=11, textColor=colors.HexColor(f"#{INK}"), leftIndent=7, rightIndent=7, spaceBefore=6, spaceAfter=6),
        "table": ParagraphStyle("TableText", fontName="MicrosoftYaHei", fontSize=7.8, leading=11, textColor=colors.HexColor(f"#{INK}")),
        "table_head": ParagraphStyle("TableHead", fontName="MicrosoftYaHei-Bold", fontSize=7.8, leading=11, textColor=colors.HexColor(f"#{INK}")),
        "callout": ParagraphStyle("Callout", fontName="MicrosoftYaHei", fontSize=9, leading=14, textColor=colors.HexColor(f"#{HEADING_DARK}")),
    }


def rl_inline(text):
    output = []
    position = 0
    for match in INLINE.finditer(text):
        output.append(escape(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            output.append(f"<b>{escape(token[2:-2])}</b>")
        elif token.startswith("`"):
            output.append(f'<font name="Consolas" color="#{HEADING_DARK}">{escape(token[1:-1])}</font>')
        else:
            label, url = re.match(r"\[(.+?)\]\((.+?)\)", token).groups()
            output.append(f'<link href="{escape(url)}" color="#{HEADING}"><b>{escape(label)}</b></link>')
        position = match.end()
    output.append(escape(text[position:]))
    return "".join(output)


def pdf_table(rows, styles):
    widths_dxa = column_widths(rows)
    data = []
    for row_index, row in enumerate(rows):
        style = styles["table_head"] if row_index == 0 else styles["table"]
        data.append([Paragraph(rl_inline(value), style) for value in row])
    table = Table(data, colWidths=[value / 1440 * inch for value in widths_dxa], repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{TABLE_FILL}")),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{LINE}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
    ]))
    return [table, Spacer(1, 6)]


def pdf_callout(text, styles):
    table = Table([[Paragraph(rl_inline(text), styles["callout"])]], colWidths=[6.5 * inch], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{CALLOUT_FILL}")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(f"#{LINE}")),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    return [table, Spacer(1, 6)]


def markdown_pdf_story(markdown_path, styles):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    first_break = lines.index("<!-- pagebreak -->")
    lines = lines[first_break + 1:]
    story = []
    index = 0
    paragraph_buffer = []
    list_kind = None
    list_items = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            story.append(Paragraph(rl_inline(" ".join(item.strip() for item in paragraph_buffer)), styles["body"]))
            paragraph_buffer = []

    def flush_list():
        nonlocal list_kind, list_items
        if list_items:
            flowable_items = [ListItem(Paragraph(rl_inline(item), styles["body"]), leftIndent=0) for item in list_items]
            options = {"bulletType": "1", "start": "1"} if list_kind == "decimal" else {"bulletType": "bullet"}
            story.append(ListFlowable(flowable_items, leftIndent=20, bulletFontName="MicrosoftYaHei", bulletFontSize=8.5, spaceAfter=6, **options))
            list_kind = None
            list_items = []

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            flush_paragraph()
            flush_list()
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            flush_paragraph(); flush_list(); story.append(PageBreak()); index += 1; continue
        if stripped.startswith("```"):
            flush_paragraph(); flush_list(); code = []; index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index]); index += 1
            code_table = Table([[XPreformatted(escape("\n".join(code)), styles["code"])]], colWidths=[6.5 * inch])
            code_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{CODE_FILL}")), ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
            story.extend([code_table, Spacer(1, 6)]); index += 1; continue
        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph(); flush_list(); image_path = (markdown_path.parent / image_match.group(2)).resolve()
            picture = PdfImage(str(image_path), width=6.3 * inch, height=4.2 * inch)
            picture.hAlign = "CENTER"
            story.append(KeepTogether([picture, Paragraph(f"图：{escape(image_match.group(1))}", styles["caption"])])); index += 1; continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.fullmatch(r"\|?[\s:|-]+\|?", lines[index + 1].strip()):
            flush_paragraph(); flush_list(); rows = [[cell.strip() for cell in stripped.strip("|").split("|")]]; index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")]); index += 1
            story.extend(pdf_table(rows, styles)); continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph(); flush_list(); level = min(3, len(heading.group(1))); story.append(Paragraph(rl_inline(heading.group(2)), styles[f"h{level}"])); index += 1; continue
        if stripped.startswith("> "):
            flush_paragraph(); flush_list(); story.extend(pdf_callout(stripped[2:], styles)); index += 1; continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or ordered:
            flush_paragraph(); kind = "bullet" if bullet else "decimal"
            if list_kind and list_kind != kind: flush_list()
            list_kind = kind; list_items.append((bullet or ordered).group(1)); index += 1; continue
        flush_list(); paragraph_buffer.append(stripped); index += 1
    flush_paragraph(); flush_list()
    return story


def build_pdf(markdown_path, output_path):
    register_pdf_fonts()
    styles = pdf_styles()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(str(output_path), pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch, title="IELTS Pilot v0.9 产品功能说明书", author="XiaoZheBrother")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")

    def decorate(canvas, document):
        canvas.saveState()
        if document.page > 1:
            canvas.setFont("MicrosoftYaHei-Bold", 7.5)
            canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
            canvas.drawString(inch, 10.55 * inch, "IELTS PILOT  /  PRODUCT GUIDE")
        canvas.setFont("MicrosoftYaHei", 7.5)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawRightString(7.5 * inch, 0.52 * inch, f"IELTS Pilot v0.9  |  {document.page}")
        canvas.restoreState()

    doc.addPageTemplates(PageTemplate(id="guide", frames=[frame], onPage=decorate))
    story = [
        Spacer(1, 0.9 * inch), Paragraph("PRODUCT GUIDE  /  RELEASE 0.9", styles["cover_kicker"]),
        Paragraph("IELTS Pilot v0.9", styles["cover_title"]), Paragraph("产品功能说明书", styles["cover_subtitle"]),
        Paragraph("本地优先的 IELTS 阅读与 AI 写作练习工作台", styles["cover_deck"]),
        Paragraph("版本 0.9.0", styles["cover_meta"]), Paragraph("2026-08-12", styles["cover_meta"]),
        Paragraph("浏览器版 · Windows 桌面版 · NSIS 安装包", styles["cover_meta"]),
        Spacer(1, 0.55 * inch), Paragraph("非官方 IELTS 产品 · 内置材料为项目原创 · 所有 Band 仅用于学习反馈", styles["cover_meta"]),
        PageBreak(),
    ]
    story.extend(markdown_pdf_story(markdown_path, styles))
    doc.build(story)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", default="docs/IELTS-Pilot-v0.9-产品功能说明书.md")
    parser.add_argument("--output", default="artifacts/manual/IELTS-Pilot-v0.9-产品功能说明书.docx")
    parser.add_argument("--pdf-output", default="artifacts/manual/IELTS-Pilot-v0.9-产品功能说明书.pdf")
    args = parser.parse_args()
    source = Path(args.source).resolve()
    output = Path(args.output).resolve()
    pdf_output = Path(args.pdf_output).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    doc.add_page_break()
    parse_manual(doc, source)
    doc.save(output)
    build_pdf(source, pdf_output)
    print(output)
    print(pdf_output)


if __name__ == "__main__":
    main()
